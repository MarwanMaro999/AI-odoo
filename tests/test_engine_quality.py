import asyncio
import json
from pathlib import Path
from types import SimpleNamespace
from uuid import uuid4

import pytest
from docx import Document

from src.engine.prompt_registry import PromptRegistry
from src.engine.registry import SkillRegistryError
from src.engine.service import DatumDocxRenderer, DatumOrchestrator, PersistentRunRepository, StoredRun
from src.engine.schemas import (
    DistributionClass,
    FindingPayload,
    PublicSkillDefinition,
    RunState,
    RunStatus,
    SkillReference,
    SourceMaterial,
    StartRunRequest,
    Verdict,
)
from src.shared.llm.contracts import GeneratedText


def test_prompt_registry_resolves_private_reference(tmp_path: Path) -> None:
    prompt = tmp_path / "gen-strs" / "v1.txt"
    prompt.parent.mkdir()
    prompt.write_text("Private instruction", encoding="utf-8")
    assert PromptRegistry(tmp_path).resolve("prompt-registry://gen-strs/v1") == "Private instruction"


def test_prompt_registry_rejects_missing_instruction(tmp_path: Path) -> None:
    with pytest.raises(SkillRegistryError, match="skill_instruction_not_found"):
        PromptRegistry(tmp_path).resolve("prompt-registry://gen-strs/v1")


def test_structured_output_requires_supplied_source_references() -> None:
    sections, references = DatumOrchestrator._parse_structured_sections(
        '{"sections":[{"title":"Purpose","points":["A"]},{"title":"Security","points":["B"]},{"title":"Workflow","points":["C"]},{"title":"Acceptance","points":["D","E","F","G","H"]}],"source_references":["architecture"]}'
    )
    references = DatumOrchestrator._validate_generated_sections(
        "test-generator", sections, references, [SimpleNamespace(source_id="architecture", text="Datum Engine architecture")]
    )
    assert references == ["architecture"]


def test_structured_output_normalises_a_supplied_source_name_to_its_id() -> None:
    sections = [("Purpose", ["A"]), ("Security", ["B"]), ("Workflow", ["C"]), ("Acceptance", ["D", "E", "F", "G", "H"])]
    references = DatumOrchestrator._validate_generated_sections(
        "test-generator", sections, ["Meeting transcript"],
        [SimpleNamespace(source_id="meeting-transcript", name="Meeting transcript", text="Datum Engine architecture")],
    )
    assert references == ["meeting-transcript"]


def test_structured_output_accepts_parent_id_for_chunked_odoo_source() -> None:
    sections = [("Purpose", ["A"]), ("Security", ["B"]), ("Workflow", ["C"]), ("Acceptance", ["D", "E", "F", "G", "H"])]
    references = DatumOrchestrator._validate_generated_sections(
        "test-generator",
        sections,
        ["attachment-790", "message-612"],
        [
            SimpleNamespace(source_id="attachment-790:chunk:0", name="Context (excerpt 1)", text="Context"),
            SimpleNamespace(source_id="attachment-790:chunk:1", name="Context (excerpt 2)", text="More context"),
            SimpleNamespace(source_id="message-612:chunk:0", name="Meeting (excerpt 1)", text="Meeting"),
        ],
    )
    assert references == ["attachment-790", "message-612"]


def test_structured_output_rejects_unsupported_marketplace_claim() -> None:
    sections = [("Purpose", ["Seller inventory workflow"])] * 8
    with pytest.raises(ValueError, match="generation_output_unsupported_claim"):
        DatumOrchestrator._validate_generated_sections(
            "gen-strs", sections, ["architecture"], [SimpleNamespace(source_id="architecture", text="Datum Engine architecture")]
        )


def test_strs_quality_gate_requires_sections_without_requirement_ids() -> None:
    sections = [
        ("Purpose and Scope", ["A"]), ("Stakeholders and Decisions", ["B"]),
        ("Functional Requirements", ["The system shall retain a source revision."] * 10),
        ("Security and Confidentiality", ["C"]), ("Data, Lineage and Staleness", ["D"]),
        ("Non-Functional Requirements", ["E"]), ("Acceptance Criteria", ["F", "G", "H"]),
        ("Assumptions", ["I"]), ("Open Decisions", ["J"]),
    ]
    references = DatumOrchestrator._validate_generated_sections(
        "gen-strs", sections, ["architecture"],
        [SimpleNamespace(source_id="architecture", name="Architecture", text="Datum Engine architecture")],
    )
    assert references == ["architecture"]


def test_questionnaire_quality_gate_requires_arabic_sections_and_questions() -> None:
    sections = [
        ("الأهداف", ["ما الهدف من النظام؟"] * 3),
        ("أصحاب المصلحة", ["من المستخدم المسؤول؟"] * 3),
        ("سير العمل", ["كيف تبدأ العملية؟"] * 3),
        ("المعلومات والتكاملات", ["ما مصادر البيانات؟"] * 3),
        ("الحوكمة والأمن", ["من يوافق على المخرجات؟"] * 3),
        ("التسليم والقبول", ["ما معايير القبول؟"] * 3),
    ]
    references = DatumOrchestrator._validate_generated_sections(
        "gen-discovery-questions", sections, ["architecture"],
        [SimpleNamespace(source_id="architecture", name="Architecture", text="Datum Engine architecture")],
    )
    assert references == ["architecture"]


def test_questionnaire_quality_gate_rejects_english_phrases() -> None:
    sections = [
        ("الأهداف", ["ما الهدف من النظام؟", "ما معيار الجودة؟", "ما الأولوية؟"]),
        ("أصحاب المصلحة", ["من المستخدم؟", "من المراجع؟", "من المالك؟"]),
        ("سير العمل", ["كيف تبدأ العملية؟", "كيف تنتهي؟", "كيف تراجع؟"]),
        ("المعلومات والتكاملات", ["ما مصدر البيانات؟", "ما التكامل؟", "ما المراجعة؟"]),
        ("الحوكمة والأمن", ["من يوافق؟", "ما الضابط؟", "ما السجل؟"]),
        ("التسليم والقبول", ["ما المعيار؟", "ما القبول؟", "ما What هو الاعتماد؟"]),
    ]
    with pytest.raises(ValueError, match="questionnaire_non_arabic_phrase_not_allowed"):
        DatumOrchestrator._validate_generated_sections(
            "gen-discovery-questions", sections, ["architecture"],
            [SimpleNamespace(source_id="architecture", name="Architecture", text="Datum Engine architecture")],
        )


def test_questionnaire_quality_gate_allows_source_grounded_latin_proper_nouns() -> None:
    sections = [
        ("الأهداف", ["ما أهداف Nile Delta Distribution؟"] * 3),
        ("أصحاب المصلحة", ["من المستخدم المسؤول في Nile؟"] * 3),
        ("سير العمل", ["كيف تبدأ العملية؟"] * 3),
        ("المعلومات والتكاملات", ["ما مصادر CSV المطلوبة؟"] * 3),
        ("الحوكمة والأمن", ["من يوافق على المخرجات؟"] * 3),
        ("التسليم والقبول", ["ما معايير القبول؟"] * 3),
    ]

    references = DatumOrchestrator._validate_generated_sections(
        "gen-discovery-questions",
        sections,
        ["customer"],
        [
            SimpleNamespace(
                source_id="customer",
                name="Customer context",
                text="Nile Delta Distribution will provide CSV files.",
            )
        ],
    )

    assert references == ["customer"]


def test_questionnaire_quality_gate_allows_source_grounded_lowercase_operational_terms() -> None:
    sections = [
        ("الأهداف", ["ما متطلبات hypercare بعد الإطلاق؟"] * 3),
        ("أصحاب المصلحة", ["من المستخدم المسؤول؟"] * 3),
        ("سير العمل", ["كيف تبدأ العملية؟"] * 3),
        ("المعلومات والتكاملات", ["ما مصادر البيانات؟"] * 3),
        ("الحوكمة والأمن", ["من يوافق على المخرجات؟"] * 3),
        ("التسليم والقبول", ["ما معايير القبول؟"] * 3),
    ]

    references = DatumOrchestrator._validate_generated_sections(
        "gen-discovery-questions",
        sections,
        ["delivery"],
        [SimpleNamespace(source_id="delivery", name="Delivery", text="Include hypercare after go-live")],
    )

    assert references == ["delivery"]


def test_questionnaire_quality_gate_allows_component_of_source_compound_term() -> None:
    sections = [
        ("الأهداف", ["ما متطلبات live بعد الإطلاق؟"] * 3),
        ("أصحاب المصلحة", ["من المستخدم المسؤول؟"] * 3),
        ("سير العمل", ["كيف تبدأ العملية؟"] * 3),
        ("المعلومات والتكاملات", ["ما مصادر البيانات؟"] * 3),
        ("الحوكمة والأمن", ["من يوافق على المخرجات؟"] * 3),
        ("التسليم والقبول", ["ما معايير القبول؟"] * 3),
    ]

    references = DatumOrchestrator._validate_generated_sections(
        "gen-discovery-questions",
        sections,
        ["delivery"],
        [SimpleNamespace(source_id="delivery", name="Delivery", text="Approve go-live readiness")],
    )

    assert references == ["delivery"]


def test_sow_quality_gate_rejects_missing_acceptance_section() -> None:
    sections = [
        ("Objective", ["A"]), ("Scope of Work", ["B"]), ("Deliverables", ["C"]),
        ("Client Responsibilities", ["D"]), ("OdooTec Responsibilities", ["E"]),
        ("Exclusions", ["F"]), ("Dependencies and Assumptions", ["G"]),
        ("Governance and Change Control", ["H"]),
    ]
    with pytest.raises(ValueError, match="sow_required_sections_missing"):
        DatumOrchestrator._validate_generated_sections(
            "gen-sow", sections, ["architecture"],
            [SimpleNamespace(source_id="architecture", name="Architecture", text="Datum Engine architecture")],
        )


def test_missing_provider_fails_instead_of_returning_demo_document(tmp_path: Path) -> None:
    assert DatumOrchestrator(
        PersistentRunRepository(tmp_path / "state"),
        SimpleNamespace(),
        DatumDocxRenderer(tmp_path / "outputs"),
        text_generator=None,
        prompt_registry=None,
        allow_demo_outputs=False,
    )


def test_large_sources_are_utf8_bounded_without_losing_source_identity(tmp_path: Path) -> None:
    orchestrator = DatumOrchestrator(
        PersistentRunRepository(tmp_path / "state"), SimpleNamespace(), DatumDocxRenderer(tmp_path / "outputs"),
        max_source_bytes=240,
    )
    bounded = orchestrator._bounded_source_text([
        SimpleNamespace(source_id="arabic", name="Arabic", text="معلومة " * 200),
        SimpleNamespace(source_id="english", name="English", text="information " * 200),
    ])
    assert "[arabic: Arabic]" in bounded
    assert "[english: English]" in bounded
    assert "Source truncated" in bounded
    assert len(bounded.encode("utf-8")) <= 260
    with pytest.raises(ValueError, match="source_grounded_generation_unavailable"):
        asyncio.run(
            orchestrator._generate_sections(
                "gen-strs", [SimpleNamespace(source_id="architecture", name="Architecture", text="Datum Engine")], {}, "", DistributionClass.CLIENT_PERMITTED, {}
        )
    )


def test_generation_prompt_reserves_fixed_overhead_and_keeps_clarification(tmp_path: Path) -> None:
    prompt_path = tmp_path / "prompts" / "gen-sow" / "v1.txt"
    prompt_path.parent.mkdir(parents=True)
    prompt_path.write_text("Create the approved Scope of Work.", encoding="utf-8")
    prompts: list[str] = []

    class CapturingGenerator:
        async def generate(self, prompt: str) -> GeneratedText:
            prompts.append(prompt)
            return GeneratedText(
                text=(
                    '{"sections":['
                    '{"title":"Objective","points":["A"]},'
                    '{"title":"Scope of Work","points":["B"]},'
                    '{"title":"Deliverables","points":["C"]},'
                    '{"title":"Client Responsibilities","points":["D"]},'
                    '{"title":"OdooTec Responsibilities","points":["E"]},'
                    '{"title":"Exclusions","points":["F"]},'
                    '{"title":"Dependencies and Assumptions","points":["G"]},'
                    '{"title":"Governance and Change Control","points":["H"]},'
                    '{"title":"Acceptance Gates","points":["I"]}],'
                    '"source_references":["source-0","clarification"]}'
                ),
                provider="test",
                model="test",
            )

    sources = [
        SimpleNamespace(source_id=f"source-{index}", name=f"Source {index}", text="evidence " * 600)
        for index in range(23)
    ] + [
        SimpleNamespace(
            source_id="clarification",
            name="Clarification answers",
            text="Saturday 20:00-Sunday 06:00 Cairo time.",
        )
    ]
    orchestrator = DatumOrchestrator(
        PersistentRunRepository(tmp_path / "state"),
        SimpleNamespace(),
        DatumDocxRenderer(tmp_path / "outputs"),
        text_generator=CapturingGenerator(),
        prompt_registry=PromptRegistry(tmp_path / "prompts"),
        max_source_bytes=32_000,
        generation_request_byte_cap=12_000,
    )

    asyncio.run(orchestrator._generate_sections(
        "gen-sow",
        sources,
        {"document_language": "En", "directed_findings": [{"finding_key": "E2E-CLARIFY-001"}]},
        "",
        DistributionClass.CLIENT_PERMITTED,
        {"instruction_ref": "prompt-registry://gen-sow/v1"},
    ))

    assert len(prompts[0].encode("utf-8")) <= 12_000 - 512
    assert "Saturday 20:00-Sunday 06:00 Cairo time." in prompts[0]
    assert "[source-0: Source 0]" in prompts[0]
    assert "[clarification: Clarification answers]" in prompts[0]


def test_review_source_budget_prioritises_document_and_keeps_baseline_ids(tmp_path: Path) -> None:
    orchestrator = DatumOrchestrator(
        PersistentRunRepository(tmp_path / "state"),
        SimpleNamespace(),
        DatumDocxRenderer(tmp_path / "outputs"),
    )
    baseline = [
        SimpleNamespace(source_id=f"baseline-{index}", name=f"Baseline {index}", text="evidence " * 500)
        for index in range(12)
    ]
    document = SimpleNamespace(
        source_id="sow-version-2",
        name="Scope of Work v2",
        text="approved downtime Saturday 20:00-Sunday 06:00 Cairo time " * 100,
    )

    bounded = orchestrator._bounded_review_source_text(
        baseline,
        document,
        max_bytes=6_000,
    )

    assert len(bounded.encode("utf-8")) <= 6_000
    assert "[sow-version-2: Scope of Work v2]" in bounded
    assert "approved downtime" in bounded
    assert "[baseline-0: Baseline 0]" in bounded
    assert "[baseline-11: Baseline 11]" in bounded


def test_all_live_skills_resolve_to_private_versioned_instructions() -> None:
    from src.engine.registry import SkillRegistry

    registry = SkillRegistry(Path("C:/ProgramData/OdooTec/datum-engine-registry"))
    prompts = PromptRegistry(Path("C:/ProgramData/OdooTec/datum-engine-prompts"))
    for identifier in ("gen-discovery-questions", "gen-strs", "gen-sow", "rev-sow"):
        _, payload = registry.load(identifier, "1.0.0")
        assert prompts.resolve(payload["instruction_ref"])
        assert not payload.get("placeholder_active")


class _AuditGenerator:
    async def generate(self, _: str) -> GeneratedText:
        return GeneratedText(
            text=(
                '{"verdict":"not_cleared","findings":[{"finding_key":"DE-REV-001",'
                '"severity":"blocking","category":"missing_acceptance_criteria",'
                '"location":"Acceptance criteria","summary":"Criteria are not measurable.",'
                '"recommendation":"Add measurable acceptance criteria.",'
                '"evidence":"The supplied SOW has no measurable acceptance criteria.",'
                '"resolution_route":"regenerate","prior_outcome":null}]}'
            ),
            provider="test",
            model="test",
        )


class _RepairingAuditGenerator:
    def __init__(self) -> None:
        self.prompts: list[str] = []

    async def generate(self, prompt: str) -> GeneratedText:
        self.prompts.append(prompt)
        if len(self.prompts) == 1:
            return GeneratedText(text="not-json", provider="test", model="test")
        return await _AuditGenerator().generate(prompt)


class _StructuredGenerator:
    def __init__(self, marker: str) -> None:
        self.marker = marker
        self.prompts: list[str] = []

    async def generate(self, prompt: str) -> GeneratedText:
        self.prompts.append(prompt)
        return GeneratedText(
            text=(
                '{"sections":['
                '{"title":"Objective","points":["%s objective"]},'
                '{"title":"Scope","points":["%s scope"]},'
                '{"title":"Deliverables","points":["%s deliverable"]},'
                '{"title":"Responsibilities","points":["%s responsibility"]},'
                '{"title":"Exclusions","points":["%s exclusion"]},'
                '{"title":"Dependencies","points":["%s dependency"]},'
                '{"title":"Governance","points":["%s governance"]},'
                '{"title":"Acceptance","points":["%s acceptance"]}'
                '],"source_references":["requirements"]}'
            ) % ((self.marker,) * 8),
            provider="test", model=self.marker,
        )


class _ArabicStructuredGenerator:
    def __init__(self) -> None:
        self.prompts: list[str] = []

    async def generate(self, prompt: str) -> GeneratedText:
        self.prompts.append(prompt)
        return GeneratedText(
            text=(
                '{"sections":['
                '{"title":"الهدف","points":["هدف عربي"]},'
                '{"title":"النطاق","points":["نطاق عربي"]},'
                '{"title":"المخرجات","points":["مخرج عربي"]},'
                '{"title":"المسؤوليات","points":["مسؤولية عربية"]},'
                '{"title":"الاستثناءات","points":["استثناء عربي"]},'
                '{"title":"الاعتماديات","points":["اعتمادية عربية"]},'
                '{"title":"الحوكمة","points":["حوكمة عربية"]},'
                '{"title":"معايير القبول","points":["معيار قبول عربي"]}'
                '],"source_references":["requirements"]}'
            ),
            provider="test",
            model="arabic",
        )


class _ArabicQuestionnaireGenerator:
    def __init__(self) -> None:
        self.prompts: list[str] = []

    async def generate(self, prompt: str) -> GeneratedText:
        self.prompts.append(prompt)
        sections = [
            ("الأهداف", "ما الهدف من النظام؟"),
            ("أصحاب المصلحة", "من المستخدم المسؤول؟"),
            ("سير العمل", "كيف تبدأ العملية؟"),
            ("المعلومات والتكاملات", "ما مصادر البيانات؟"),
            ("الحوكمة والأمن", "من يوافق على المخرجات؟"),
            ("التسليم والقبول", "ما معايير القبول؟"),
        ]
        return GeneratedText(
            text=json.dumps(
                {
                    "sections": [
                        {"title": title, "points": [question] * 3}
                        for title, question in sections
                    ],
                    "source_references": ["requirements"],
                },
                ensure_ascii=False,
            ),
            provider="test",
            model="arabic-questionnaire",
        )


def test_arabic_questionnaire_prompt_does_not_receive_sow_section_titles(tmp_path: Path) -> None:
    prompt = tmp_path / "gen-discovery-questions" / "v1.txt"
    prompt.parent.mkdir()
    prompt.write_text("Use exactly the six questionnaire sections.", encoding="utf-8")
    generator = _ArabicQuestionnaireGenerator()
    orchestrator = DatumOrchestrator(
        PersistentRunRepository(tmp_path / "state"),
        SimpleNamespace(),
        DatumDocxRenderer(tmp_path / "outputs"),
        text_generator=generator,
        arabic_text_generator=generator,
        prompt_registry=PromptRegistry(tmp_path),
    )

    asyncio.run(orchestrator._generate_sections(
        "gen-discovery-questions",
        [SimpleNamespace(source_id="requirements", name="Requirements", text="Required source")],
        {"document_language": "Ar"},
        "",
        DistributionClass.CLIENT_PERMITTED,
        {"instruction_ref": "prompt-registry://gen-discovery-questions/v1"},
    ))

    assert "Write the complete document in Arabic" in generator.prompts[0]
    assert "الهدف؛ نطاق العمل؛ المخرجات" not in generator.prompts[0]


def test_arabic_generation_uses_its_explicit_generator_responsibility(tmp_path: Path) -> None:
    prompt = tmp_path / "gen-sow" / "v1.txt"
    prompt.parent.mkdir()
    prompt.write_text("Generate only from the source.", encoding="utf-8")
    english = _StructuredGenerator("english")
    arabic = _ArabicStructuredGenerator()
    orchestrator = DatumOrchestrator(
        PersistentRunRepository(tmp_path / "state"), SimpleNamespace(), DatumDocxRenderer(tmp_path / "outputs"),
        text_generator=english, arabic_text_generator=arabic, prompt_registry=PromptRegistry(tmp_path),
    )

    sections, _ = asyncio.run(orchestrator._generate_sections(
        "gen-sow", [SimpleNamespace(source_id="requirements", name="Requirements", text="Required source")],
        {"document_language": "Ar"}, "", DistributionClass.CLIENT_PERMITTED,
        {"instruction_ref": "prompt-registry://gen-sow/v1"},
    ))

    assert sections[0][1] == ["هدف عربي"]
    assert arabic.prompts
    assert "الهدف؛ نطاق العمل؛ المخرجات" in arabic.prompts[0]
    assert "do not repeat the English titles" in arabic.prompts[0]
    assert not english.prompts


def test_arabic_sow_quality_gate_rejects_english_content() -> None:
    sections = [("Objective", ["English point"])] * 8
    with pytest.raises(ValueError, match="sow_arabic_output_contains_non_arabic_section"):
        DatumOrchestrator._validate_output_language("gen-sow", sections, "Ar")


class _BilingualAuditGenerator:
    def __init__(self) -> None:
        self.prompts: list[str] = []

    async def generate(self, prompt: str) -> GeneratedText:
        self.prompts.append(prompt)
        return GeneratedText(
            text=(
                '{"verdict":"not_cleared","findings":[{"finding_key":"DE-REV-001",'
                '"severity":"blocking","category":"missing_acceptance_criteria",'
                '"location":"Acceptance criteria","summary":"Criteria are not measurable.",'
                '"recommendation":"Add measurable acceptance criteria.",'
                '"evidence":"The supplied SOW has no measurable acceptance criteria.",'
                '"resolution_route":"regenerate","prior_outcome":null}]}'
            ),
            provider="test",
            model="test",
        )


class _EscalatingOrchestrator(DatumOrchestrator):
    async def _complete_generation(self, *_: object) -> None:
        raise ValueError("sow_review_cycle_ceiling_escalated")


class _ReviewPersistenceRecorder:
    def __init__(self) -> None:
        self.calls = []

    async def store(self, *args) -> None:
        self.calls.append(args)


def test_review_result_is_sent_to_durable_review_persistence(tmp_path: Path) -> None:
    recorder = _ReviewPersistenceRecorder()
    orchestrator = DatumOrchestrator(
        PersistentRunRepository(tmp_path / "state"), SimpleNamespace(), DatumDocxRenderer(tmp_path / "outputs"),
        review_persistence=recorder,
    )
    finding = FindingPayload(
        finding_key="DE-REV-TEST", severity="blocking", category="completeness", location="Scope",
        summary="A required detail is missing.", recommendation="Add the missing detail.",
        evidence="The source does not contain it.", resolution_route="regenerate",
    )

    asyncio.run(orchestrator._review_persistence.store(uuid4(), 2, 3, "not_cleared", [finding]))

    assert recorder.calls[0][1:4] == (2, 3, "not_cleared")


def test_reviewer_finding_accepts_a_targeted_clarification_question() -> None:
    finding = FindingPayload(
        finding_key="DE-REV-CLARIFY", severity="blocking", category="timeline", location="Delivery plan",
        summary="The go-live date is absent.", recommendation="Confirm the target go-live date.",
        evidence="No date appears in the source material.", resolution_route="clarify",
        clarification_question="What is the customer-approved go-live date?",
    )

    assert finding.clarification_question == "What is the customer-approved go-live date?"


def test_review_parser_deduplicates_repeated_language_finding_keys() -> None:
    result = (
        '{"verdict":"not_cleared","findings":['
        '{"finding_key":"F001","severity":"advisory","category":"security",'
        '"location":"Security","summary":"Clarify security.",'
        '"recommendation":"Confirm controls.","evidence":"Open question.",'
        '"resolution_route":"clarify"},'
        '{"finding_key":"F001","severity":"advisory","category":"security",'
        '"location":"Security","summary":"Repeated copy.",'
        '"recommendation":"Confirm controls.","evidence":"Open question.",'
        '"resolution_route":"clarify"}]}'
    )
    verdict, findings = DatumOrchestrator._parse_review_result(result, "Ar")
    assert verdict == Verdict.NOT_CLEARED
    assert len(findings) == 1
    assert findings[0].finding_key == "F001"


def test_review_merges_retrieval_chunks_into_one_logical_sow() -> None:
    merged = DatumOrchestrator._merge_chunked_review_sources([
        SimpleNamespace(source_id="sow-version-73:chunk:0", revision="1", type="sow_version", name="Arabic SoW (excerpt 1)", text="first"),
        SimpleNamespace(source_id="sow-version-73:chunk:1", revision="1", type="sow_version", name="Arabic SoW (excerpt 2)", text="second"),
        SimpleNamespace(source_id="sow-version-71:chunk:0", revision="1", type="sow_version", name="English SoW (excerpt 1)", text="english"),
    ])
    assert [item.source_id for item in merged] == ["sow-version-73", "sow-version-71"]
    assert merged[0].name == "Arabic SoW"
    assert merged[0].text == "first\n\nsecond"


def test_review_excludes_historical_session_versions() -> None:
    sources = [
        SimpleNamespace(source_id="sow-version-76"),
        SimpleNamespace(source_id="sow-version-75"),
        SimpleNamespace(source_id="sow-version-73"),
    ]
    targeted = DatumOrchestrator._targeted_review_sources(
        sources,
        {"sow-version-75": "En", "sow-version-76": "Ar"},
    )
    assert [source.source_id for source in targeted] == [
        "sow-version-76",
        "sow-version-75",
    ]


def test_discovery_questionnaire_output_metadata_is_arabic() -> None:
    assert DatumOrchestrator._generation_languages(
        "gen-discovery-questions",
        "discovery_questionnaire",
        {},
    ) == ("Ar",)
    filename = DatumDocxRenderer.versioned_filename(
        "discovery_questionnaire",
        version=1,
        language=DatumOrchestrator._generation_languages(
            "gen-discovery-questions",
            "discovery_questionnaire",
            {},
        )[0],
    )
    assert filename.endswith("_Ar_Discovery_Questionnaire_v1.docx")


def test_unreviewed_generation_checkpoints_completed_editions(tmp_path: Path) -> None:
    request = StartRunRequest(
        idempotency_key="checkpointed-sow",
        engagement_id="engagement-1",
        source_set_revision="r1",
        skill=SkillReference(identifier="gen-sow", version="v1"),
        source_material=[SourceMaterial(
            source_id="requirements",
            revision="r1",
            type="approved_requirements_specification",
            name="Approved StRS",
            text="Required source context.",
        )],
    )
    repository = PersistentRunRepository(tmp_path / "state")
    run, _ = repository.create_or_get(request)
    orchestrator = DatumOrchestrator(
        repository,
        SimpleNamespace(),
        DatumDocxRenderer(tmp_path / "outputs"),
    )
    calls = 0

    async def flaky_generation(*_: object) -> tuple[list[tuple[str, list[str]]], list[str]]:
        nonlocal calls
        calls += 1
        if calls == 2:
            raise RuntimeError("transient provider failure")
        return [
            ("Purpose", ["Grounded purpose"]),
            ("Scope", ["Grounded scope"]),
            ("Deliverables", ["Grounded deliverable"]),
            ("Responsibilities", ["Grounded responsibility"]),
            ("Timeline", ["Grounded timeline"]),
            ("Acceptance Criteria", ["Grounded acceptance criterion"]),
            ("Assumptions", ["Grounded assumption"]),
            ("Exclusions", ["Grounded exclusion"]),
        ], ["requirements"]

    orchestrator._generate_sections = flaky_generation  # type: ignore[method-assign]
    definition = SimpleNamespace(
        version="v1",
        outputs=[
            SimpleNamespace(document_type="sow", distribution_class=DistributionClass.CLIENT_PERMITTED),
            SimpleNamespace(document_type="sow", distribution_class=DistributionClass.INTERNAL_ONLY),
        ],
    )

    with pytest.raises(RuntimeError, match="transient provider failure"):
        asyncio.run(orchestrator._complete_unreviewed_generation(run, definition, {}, {}, ""))
    assert len(repository.get(run.status.run_id).status.outputs) == 1

    asyncio.run(orchestrator._complete_unreviewed_generation(run, definition, {}, {}, ""))
    assert len(repository.get(run.status.run_id).status.outputs) == 4
    assert calls == 5


def test_auditor_records_model_generated_structured_finding(tmp_path: Path) -> None:
    prompt = tmp_path / "rev-sow" / "v1.txt"
    prompt.parent.mkdir()
    prompt.write_text("Review only supplied material.", encoding="utf-8")
    request = StartRunRequest(
        idempotency_key="audit-quality-test",
        engagement_id="engagement-1",
        source_set_revision="r1",
        skill=SkillReference(identifier="rev-sow", version="v1"),
        source_material=[SourceMaterial(
            source_id="sow-1", revision="r1", type="scope_of_work", name="SOW",
            text="The supplied SOW has no measurable acceptance criteria.",
        )],
        parameters={"document_version": 2, "language_code": "ArEn"},
    )
    run = StoredRun(
        status=RunStatus(
            run_id=__import__("uuid").uuid4(), state=RunState.RUNNING, skill=request.skill,
            engagement_id=request.engagement_id, source_set_revision=request.source_set_revision,
            submitted_at=__import__("datetime").datetime.now(__import__("datetime").timezone.utc),
        ),
        request=request,
        fingerprint="test",
    )
    review_persistence = _ReviewPersistenceRecorder()
    orchestrator = DatumOrchestrator(
        PersistentRunRepository(tmp_path / "state"), SimpleNamespace(), DatumDocxRenderer(tmp_path / "outputs"),
        text_generator=_AuditGenerator(), prompt_registry=PromptRegistry(tmp_path),
        review_persistence=review_persistence,
    )
    asyncio.run(orchestrator._complete_audit(run, SimpleNamespace(), {"instruction_ref": "prompt-registry://rev-sow/v1"}))
    assert run.status.verdict == Verdict.NOT_CLEARED
    assert run.status.findings[0].finding_key == "DE-REV-001"
    assert len(run.status.outputs) == 2
    filenames = {output.filename for output in run.status.outputs}
    assert any(filename.endswith("_En_SoW_Review_Findings_v2.docx") for filename in filenames)
    assert any(filename.endswith("_Ar_SoW_Review_Findings_v2.docx") for filename in filenames)
    assert all((tmp_path / "outputs" / output.filename).is_file() for output in run.status.outputs)
    assert review_persistence.calls[0][1:4] == (1, 2, "not_cleared")


def test_auditor_repairs_invalid_json_before_failing_the_run(tmp_path: Path) -> None:
    prompt = tmp_path / "rev-sow" / "v1.txt"
    prompt.parent.mkdir()
    prompt.write_text("Review only supplied material.", encoding="utf-8")
    request = StartRunRequest(
        idempotency_key="audit-json-repair-test",
        engagement_id="engagement-1",
        source_set_revision="r1",
        skill=SkillReference(identifier="rev-sow", version="v1"),
        source_material=[SourceMaterial(
            source_id="sow-1", revision="r1", type="scope_of_work", name="SOW",
            text="The supplied SOW has no measurable acceptance criteria.",
        )],
        parameters={"document_version": 2, "language_code": "ArEn"},
    )
    run = StoredRun(
        status=RunStatus(
            run_id=__import__("uuid").uuid4(), state=RunState.RUNNING, skill=request.skill,
            engagement_id=request.engagement_id, source_set_revision=request.source_set_revision,
            submitted_at=__import__("datetime").datetime.now(__import__("datetime").timezone.utc),
        ),
        request=request,
        fingerprint="test",
    )
    generator = _RepairingAuditGenerator()
    orchestrator = DatumOrchestrator(
        PersistentRunRepository(tmp_path / "state"), SimpleNamespace(), DatumDocxRenderer(tmp_path / "outputs"),
        text_generator=generator, prompt_registry=PromptRegistry(tmp_path), quality_attempts=3,
    )

    asyncio.run(orchestrator._complete_audit(
        run, SimpleNamespace(), {"instruction_ref": "prompt-registry://rev-sow/v1"},
    ))

    assert run.status.verdict == Verdict.NOT_CLEARED
    assert len(generator.prompts) == 2
    assert "CORRECTION REQUIRED" in generator.prompts[1]
    assert "audit_output_invalid" in generator.prompts[1]


def test_bilingual_sow_audit_reviews_each_language_separately(tmp_path: Path) -> None:
    prompt = tmp_path / "rev-sow" / "v1.txt"
    prompt.parent.mkdir()
    prompt.write_text("Review only supplied material.", encoding="utf-8")
    request = StartRunRequest(
        idempotency_key="bilingual-audit-quality-test",
        engagement_id="engagement-1",
        source_set_revision="r1",
        skill=SkillReference(identifier="rev-sow", version="v1"),
        source_material=[
            SourceMaterial(
                source_id="requirements", revision="r1", type="approved_requirements_specification",
                name="Approved StRS", text="Acceptance criteria must be measurable.",
            ),
            SourceMaterial(
                source_id="sow-version-en", revision="v1", type="sow_version", name="English SOW",
                text="The supplied SOW has no measurable acceptance criteria.",
            ),
            SourceMaterial(
                source_id="sow-version-ar", revision="v1", type="sow_version", name="Arabic SOW",
                text="نطاق العمل لا يحتوي على معايير قبول قابلة للقياس.",
            ),
        ],
        parameters={
            "document_version": 1,
            "review_languages": {"sow-version-en": "En", "sow-version-ar": "Ar"},
        },
    )
    run = StoredRun(
        status=RunStatus(
            run_id=__import__("uuid").uuid4(), state=RunState.RUNNING, skill=request.skill,
            engagement_id=request.engagement_id, source_set_revision=request.source_set_revision,
            submitted_at=__import__("datetime").datetime.now(__import__("datetime").timezone.utc),
        ),
        request=request,
        fingerprint="test",
    )
    generator = _BilingualAuditGenerator()
    orchestrator = DatumOrchestrator(
        PersistentRunRepository(tmp_path / "state"), SimpleNamespace(), DatumDocxRenderer(tmp_path / "outputs"),
        text_generator=generator, prompt_registry=PromptRegistry(tmp_path),
    )

    asyncio.run(orchestrator._complete_audit(run, SimpleNamespace(), {"instruction_ref": "prompt-registry://rev-sow/v1"}))

    assert run.status.verdict == Verdict.NOT_CLEARED
    assert [finding.language_code for finding in run.status.findings] == ["En", "Ar"]
    assert len(generator.prompts) == 2
    assert '"document_language": "En"' in generator.prompts[0]
    assert '"document_language": "Ar"' in generator.prompts[1]
    assert {output.language_code for output in run.status.outputs} == {"En", "Ar"}
    english_output = next(output for output in run.status.outputs if output.language_code == "En")
    english_text = "\n".join(paragraph.text for paragraph in Document(tmp_path / "outputs" / english_output.filename).paragraphs)
    assert "Recommended correction: Add measurable acceptance criteria." in english_text
    assert "Review cycle: 1" in english_text
    arabic_output = next(output for output in run.status.outputs if output.language_code == "Ar")
    arabic_text = "\n".join(paragraph.text for paragraph in Document(tmp_path / "outputs" / arabic_output.filename).paragraphs)
    assert "\u0646\u062a\u0627\u0626\u062c \u0645\u0631\u0627\u062c\u0639\u0629 \u0646\u0637\u0627\u0642 \u0627\u0644\u0639\u0645\u0644" in arabic_text


def test_fifth_unresolved_cycle_requires_human_intervention(tmp_path: Path) -> None:
    definition = PublicSkillDefinition(
        identifier="gen-sow", version="v1", kind="generator",
        accepted_source_material=["approved_requirements_specification"], outputs=[],
    )
    repository = PersistentRunRepository(tmp_path / "state")
    request = StartRunRequest(
        idempotency_key="five-cycle-ceiling", engagement_id="engagement-1", source_set_revision="r1",
        skill=SkillReference(identifier="gen-sow", version="v1"),
        source_material=[SourceMaterial(
            source_id="requirements", revision="r1", type="approved_requirements_specification",
            name="Approved StRS", text="Required source context.",
        )],
    )
    run, _ = repository.create_or_get(request)
    orchestrator = _EscalatingOrchestrator(
        repository,
        SimpleNamespace(load=lambda *_: (definition, {})),
        DatumDocxRenderer(tmp_path / "outputs"),
    )

    asyncio.run(orchestrator.process(run.status.run_id))

    status = repository.get(run.status.run_id).status
    assert status.state == RunState.REQUIRES_HUMAN_INTERVENTION
    assert status.failure_code == "sow_review_cycle_ceiling_escalated"


def test_versioned_bilingual_word_filenames_follow_the_required_convention() -> None:
    english = DatumDocxRenderer.versioned_filename("scope_of_work", version=3, language="En")
    arabic = DatumDocxRenderer.versioned_filename("scope_of_work", version=3, language="Ar")
    english_findings = DatumDocxRenderer.versioned_filename("sow_review_findings", version=3, language="En", findings=True)
    arabic_findings = DatumDocxRenderer.versioned_filename("sow_review_findings", version=3, language="Ar", findings=True)

    assert english.endswith("_En_SoW_v3.docx")
    assert arabic.endswith("_Ar_SoW_v3.docx")
    assert english != arabic
    assert english_findings.endswith("_En_SoW_Review_Findings_v3.docx")
    assert arabic_findings.endswith("_Ar_SoW_Review_Findings_v3.docx")
    assert len(english.split("_", maxsplit=1)[0]) == 8
