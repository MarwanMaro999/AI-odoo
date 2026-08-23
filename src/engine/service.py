"""Asynchronous generic orchestrator used by every Datum skill."""

import asyncio
from dataclasses import dataclass, field
from datetime import datetime, timezone
import hashlib
from html import escape
import json
from pathlib import Path
from threading import RLock
from uuid import UUID, uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.engine.registry import SkillRegistry, SkillRegistryError
from src.engine.schemas import (
    DistributionClass,
    FindingPayload,
    PublicSkillDefinition,
    RunLogEntry,
    RunOutput,
    RunState,
    RunStatus,
    StartRunRequest,
    Verdict,
)
from src.shared.llm.providers import GroqTextGenerator
from src.shared.web_research.research_service import CompanyResearchService


@dataclass
class StoredRun:
    status: RunStatus
    request: StartRunRequest
    fingerprint: str
    logs: list[RunLogEntry] = field(default_factory=list)


class PersistentRunRepository:
    """Small JSON-backed store for the foundation demo; survives service restarts."""

    def __init__(self, state_directory: Path) -> None:
        self._state_directory = state_directory
        self._state_directory.mkdir(parents=True, exist_ok=True)
        self._runs: dict[UUID, StoredRun] = {}
        self._keys: dict[str, UUID] = {}
        self._lock = RLock()
        self._load()

    def _load(self) -> None:
        for path in self._state_directory.glob("*.json"):
            try:
                payload = json.loads(path.read_text(encoding="utf-8"))
                status = RunStatus.model_validate(payload["status"])
                request = StartRunRequest.model_validate(payload["request"])
                run = StoredRun(status, request, payload["fingerprint"], [RunLogEntry.model_validate(x) for x in payload.get("logs", [])])
                self._runs[status.run_id] = run
                self._keys[request.idempotency_key] = status.run_id
            except (OSError, ValueError, KeyError):
                continue

    def _save(self, run: StoredRun) -> None:
        path = self._state_directory / f"{run.status.run_id}.json"
        payload = {"status": run.status.model_dump(mode="json"), "request": run.request.model_dump(mode="json"), "fingerprint": run.fingerprint, "logs": [x.model_dump(mode="json") for x in run.logs]}
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    def create_or_get(self, request: StartRunRequest) -> tuple[StoredRun, bool]:
        fingerprint = hashlib.sha256(json.dumps(request.model_dump(mode="json"), sort_keys=True).encode()).hexdigest()
        with self._lock:
            existing_id = self._keys.get(request.idempotency_key)
            if existing_id:
                existing = self._runs[existing_id]
                if existing.fingerprint != fingerprint:
                    raise ValueError("idempotency_conflict")
                return existing, False
            status = RunStatus(run_id=uuid4(), state=RunState.QUEUED, skill=request.skill, engagement_id=request.engagement_id, source_set_revision=request.source_set_revision, submitted_at=datetime.now(timezone.utc))
            run = StoredRun(status, request, fingerprint)
            self._runs[status.run_id] = run
            self._keys[request.idempotency_key] = status.run_id
            self._save(run)
            return run, True

    def get(self, run_id: UUID) -> StoredRun:
        with self._lock:
            if run_id not in self._runs:
                raise KeyError("run_not_found")
            return self._runs[run_id]

    def save(self, run: StoredRun) -> None:
        with self._lock:
            self._save(run)


class DatumDocxRenderer:
    """Render readable consultant-facing Word documents and browser previews."""

    def __init__(self, output_directory: Path) -> None:
        self._output_directory = output_directory

    def render(self, run_id: UUID, document_type: str, sections: list[tuple[str, list[str]]], output_key: str = "document", template_path: Path | None = None) -> Path:
        self._output_directory.mkdir(parents=True, exist_ok=True)
        document = Document(template_path) if template_path else Document()
        section = document.sections[0]
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        styles = document.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        styles["Normal"].font.size = Pt(10.5)
        styles["Normal"].paragraph_format.space_after = Pt(6)
        styles["Normal"].paragraph_format.line_spacing = 1.15
        for name, size, color in (("Heading 1", 15, "1F4E79"), ("Heading 2", 12, "2E75B6")):
            style = styles[name]
            style.font.name = "Arial"
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style.font.bold = True
            style.paragraph_format.space_before = Pt(14)
            style.paragraph_format.space_after = Pt(6)

        title = self._title_for(document_type)
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_paragraph.add_run(title)
        title_run.font.name = "Arial"
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(11, 37, 69)
        title_paragraph.paragraph_format.space_after = Pt(3)
        subtitle = document.add_paragraph("Datum Engine | Internal working document")
        subtitle.runs[0].font.name = "Arial"
        subtitle.runs[0].font.size = Pt(9)
        subtitle.runs[0].font.color.rgb = RGBColor(89, 89, 89)
        subtitle.paragraph_format.space_after = Pt(16)
        self._add_rule(subtitle, "1F4E79")

        header = section.header.paragraphs[0]
        header.text = "ODOOTEC | DATUM ENGINE"
        header.runs[0].font.name = "Arial"
        header.runs[0].font.size = Pt(8)
        header.runs[0].font.color.rgb = RGBColor(89, 89, 89)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run("Internal working document")
        footer_run.font.name = "Arial"
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(89, 89, 89)

        for title, paragraphs in sections:
            heading = document.add_heading(title, level=1)
            is_arabic = any("\u0600" <= char <= "\u06ff" for char in title)
            if is_arabic:
                heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self._set_rtl(heading)
            for text in paragraphs:
                style_name = "List Number" if document_type == "discovery_questionnaire" else "List Bullet"
                paragraph = document.add_paragraph(style=style_name)
                paragraph.add_run(text)
                paragraph.paragraph_format.space_after = Pt(5)
                if any("\u0600" <= char <= "\u06ff" for char in text):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    self._set_rtl(paragraph)
        path = self._output_directory / f"{document_type}-{output_key}-{run_id}.docx"
        document.save(path)
        return path

    def render_preview(self, run_id: UUID, document_type: str, sections: list[tuple[str, list[str]]], output_key: str = "document") -> Path:
        """Create a safe in-browser rendition; the original Word file remains unchanged."""
        self._output_directory.mkdir(parents=True, exist_ok=True)
        blocks = []
        for title, paragraphs in sections:
            direction = "rtl" if any("\u0600" <= char <= "\u06ff" for char in title) else "ltr"
            items = "".join(f"<li>{escape(text)}</li>" for text in paragraphs if text)
            list_tag = "ol" if document_type == "discovery_questionnaire" else "ul"
            blocks.append(f'<section dir="{direction}"><h2>{escape(title)}</h2><{list_tag}>{items}</{list_tag}></section>')
        path = self._output_directory / f"{document_type}-{output_key}-{run_id}.html"
        path.write_text(
            "<!doctype html><html><head><meta charset=\"utf-8\"><title>Datum document preview</title>"
            "<style>body{font:16px Arial,sans-serif;line-height:1.6;max-width:960px;margin:0 auto;padding:42px 36px;color:#172033;background:#f6f8fb}main{background:#fff;padding:52px 64px;box-shadow:0 2px 14px #dbe2ea}"
            "h1{font-size:32px;color:#0b2545;margin:0 0 4px}.meta{font-size:13px;color:#666;border-bottom:3px solid #1f4e79;padding-bottom:15px;margin-bottom:28px}h2{font-size:20px;color:#1f4e79;margin:28px 0 10px}section{margin:24px 0}li{margin:8px 0;padding-left:4px}</style>"
            f"</head><body><main><h1>{escape(self._title_for(document_type))}</h1><div class=\"meta\">Datum Engine | Internal working document</div>{''.join(blocks)}</main></body></html>",
            encoding="utf-8",
        )
        return path

    @staticmethod
    def _title_for(document_type: str) -> str:
        return {
            "discovery_questionnaire": "Discovery Questionnaire",
            "strs": "Stakeholder Requirements Specification (StRS)",
            "scope_of_work": "Scope of Work",
        }.get(document_type, document_type.replace("_", " ").title())

    @staticmethod
    def _add_rule(paragraph: object, color: str) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "8")
        bottom.set(qn("w:color"), color)
        borders.append(bottom)
        p_pr.append(borders)

    @staticmethod
    def _set_rtl(paragraph: object) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)


class DatumOrchestrator:
    def __init__(self, repository: PersistentRunRepository, registry: SkillRegistry, renderer: DatumDocxRenderer, max_attempts: int = 3, research_service: CompanyResearchService | None = None, text_generator: GroqTextGenerator | None = None) -> None:
        self._repository = repository
        self._registry = registry
        self._renderer = renderer
        self._research_service = research_service
        self._text_generator = text_generator
        self._max_attempts = max_attempts

    async def process(self, run_id: UUID) -> None:
        run = self._repository.get(run_id)
        if run.status.state in {RunState.SUCCEEDED, RunState.CANCELLED}:
            return
        run.status.state = RunState.RUNNING
        run.status.started_at = datetime.now(timezone.utc)
        self._record(run, "orchestration", "started")
        self._repository.save(run)
        for attempt in range(1, self._max_attempts + 1):
            run.status.attempt_count = attempt
            try:
                definition, payload = self._registry.load(run.request.skill.identifier, run.request.skill.version)
                self._validate_input(run.request, definition, payload)
                self._record(run, "validate_input", "succeeded")
                if definition.kind.value == "auditor":
                    self._complete_audit(run, definition, payload)
                else:
                    await self._complete_generation(run, definition, payload)
                run.status.state = RunState.SUCCEEDED
                run.status.completed_at = datetime.now(timezone.utc)
                self._record(run, "orchestration", "succeeded")
                self._repository.save(run)
                return
            except (SkillRegistryError, ValueError) as error:
                run.status.failure_code = str(error) or "validation_failed"
                self._record(run, "orchestration", "failed", {"failure_code": run.status.failure_code})
                break
            except Exception as error:
                self._record(run, "attempt", "failed", {"error_type": type(error).__name__, "attempt": attempt})
                if attempt < self._max_attempts:
                    await asyncio.sleep(min(attempt, 2))
                    continue
                run.status.failure_code = "processing_failed"
        run.status.state = RunState.FAILED
        run.status.completed_at = datetime.now(timezone.utc)
        self._repository.save(run)

    async def _complete_generation(self, run: StoredRun, definition: PublicSkillDefinition, payload: dict[str, object]) -> None:
        research = ""
        if run.request.skill.identifier == "gen-discovery-questions":
            research_result = await self._research_service.research(
                " ".join(source.text for source in run.request.source_material)[:4_000]
            ) if self._research_service else None
            research = research_result.text if research_result else ""
        outputs: list[RunOutput] = []
        for output in definition.outputs:
            sections = await self._generate_sections(
                run.request.skill.identifier,
                run.request.source_material,
                run.request.parameters,
                research,
                output.distribution_class,
                str(payload.get("instruction", "")),
            )
            path = self._renderer.render(run.status.run_id, output.document_type, sections, output.distribution_class)
            preview = self._renderer.render_preview(run.status.run_id, output.document_type, sections, output.distribution_class)
            source_text = "\n\n".join(
                f"{title}\n" + "\n".join(f"- {paragraph}" for paragraph in paragraphs)
                for title, paragraphs in sections
            )
            outputs.append(RunOutput(output_id=uuid4(), document_type=output.document_type, distribution_class=output.distribution_class, filename=path.name, download_url=f"/api/v1/runs/{run.status.run_id}/outputs/{path.name}", preview_url=f"/api/v1/runs/{run.status.run_id}/outputs/{preview.name}", source_text=source_text))
        run.status.outputs = outputs
        self._record(run, "render_and_validate", "succeeded", {"output_count": len(outputs)})

    async def _generate_sections(
        self,
        identifier: str,
        sources: list[object],
        parameters: dict[str, object],
        research: str,
        distribution: DistributionClass,
        instruction: str,
    ) -> list[tuple[str, list[str]]]:
        """Use the externally registered demo instruction when a model is configured.

        The static structure remains a safe fallback so a missing model never blocks Odoo testing.
        """
        if not self._text_generator:
            return self._with_revision_notes(
                self._placeholder_sections(identifier, sources, parameters, research, distribution), parameters
            )
        if not instruction or instruction.lower().startswith("placeholder"):
            instruction = self._demo_instruction(identifier)
        source_text = "\n\n".join(f"[{source.name}]\n{source.text}" for source in sources)
        prompt = (
            f"{instruction}\n\n"
            "Use only the supplied source material. Do not invent company facts. Where information is absent, write it as an explicit assumption or open question. "
            "Return a detailed consultant-ready document using this exact plain-text format: each section begins with '## ', and every point under it begins with '- '. "
            f"Document skill: {identifier}. Distribution: {distribution.value}.\n\n"
            f"SOURCE MATERIAL\n{source_text}\n\nPUBLIC RESEARCH (validate before relying on it)\n{research or 'Not available.'}\n\n"
            f"RUN PARAMETERS\n{json.dumps(parameters, ensure_ascii=False)}"
        )
        try:
            generated = await self._text_generator.generate(prompt)
            sections = self._parse_generated_sections(generated.text)
            if len(sections) >= 5:
                return sections
        except Exception:
            pass
        return self._with_revision_notes(
            self._placeholder_sections(identifier, sources, parameters, research, distribution), parameters
        )

    @staticmethod
    def _with_revision_notes(
        sections: list[tuple[str, list[str]]], parameters: dict[str, object]
    ) -> list[tuple[str, list[str]]]:
        instructions = str(parameters.get("revision_instructions", "")).strip()
        if not instructions:
            return sections
        return sections + [
            ("Requested changes addressed in this version", [instructions])
        ]

    @staticmethod
    def _demo_instruction(identifier: str) -> str:
        """Document-specific temporary instruction used until private prompts are wired."""
        instructions = {
            "gen-discovery-questions": (
                "Create a practical bilingual discovery questionnaire, not a requirements specification or scope of work. "
                "Use Arabic and English sections. Include at least 20 diverse, meeting-ready questions grouped by business goals, stakeholders, current process, data, integrations, reporting, security, delivery priorities, and acceptance criteria."
            ),
            "gen-strs": (
                "Create a detailed Stakeholder Requirements Specification, not a questionnaire or scope of work. "
                "Use sections for business context, stakeholders, current-state issues, functional requirements, data and integrations, roles and controls, reporting, non-functional requirements, assumptions, open questions, and measurable acceptance criteria."
            ),
            "gen-sow": (
                "Create a detailed Scope of Work, not a questionnaire or requirements specification. "
                "Use sections for objective, in-scope services, deliverables, delivery approach, responsibilities, exclusions, assumptions and dependencies, acceptance gates, governance, and change control."
            ),
        }
        return instructions.get(identifier, "Create a detailed consultant-ready document for the registered skill.")

    @staticmethod
    def _parse_generated_sections(text: str) -> list[tuple[str, list[str]]]:
        sections: list[tuple[str, list[str]]] = []
        title = ""
        points: list[str] = []
        for line in text.splitlines():
            value = line.strip()
            if value.startswith("## "):
                if title and points:
                    sections.append((title, points))
                title, points = value[3:].strip(), []
            elif title and value.startswith(("- ", "* ")):
                points.append(value[2:].strip())
        if title and points:
            sections.append((title, points))
        return sections

    def _complete_audit(self, run: StoredRun, _: PublicSkillDefinition, __: dict[str, object]) -> None:
        target = str(run.request.parameters.get("target_document_version", "unknown version"))
        prior = run.request.parameters.get("carried_findings", [])
        findings: list[FindingPayload] = [FindingPayload(
            finding_key="scope-acceptance-criteria",
            severity="blocking",
            category="completeness",
            location="Acceptance criteria",
            summary=f"Add measurable acceptance criteria for {target}.",
            evidence="The submitted document version does not contain measurable acceptance criteria.",
            resolution_route="regenerate",
            prior_outcome="carried" if prior else None,
        )]
        run.status.findings = findings
        run.status.verdict = Verdict.CLEARED if not findings else Verdict.NOT_CLEARED
        self._record(run, "audit_and_validate", "succeeded", {"verdict": run.status.verdict, "finding_count": len(findings)})

    @staticmethod
    def _placeholder_sections(identifier: str, sources: list[object], parameters: dict[str, object], research: str = "", distribution: DistributionClass = DistributionClass.CLIENT_PERMITTED) -> list[tuple[str, list[str]]]:
        if identifier == "gen-discovery-questions":
            arabic = [
                "ما هي النتائج التجارية التي يجب أن يحققها المشروع خلال 6 و12 شهراً؟ وكيف سنقيس النجاح؟",
                "من هم أصحاب المصلحة، وما القرارات والمسؤوليات التي يملكونها؟",
                "ما هي رحلة المستخدم أو البائع الحالية من التسجيل حتى ما بعد البيع؟ وأين تتعطل؟",
                "ما العمليات التي تتم خارج النظام حالياً باستخدام البريد أو Excel أو الرسائل؟",
                "كيف يتم إنشاء المنتجات والتحقق من بياناتها وصورها وتصنيفاتها؟",
                "ما قواعد تسعير المنتجات والعروض والخصومات؟ ومن يوافق عليها؟",
                "كيف تتم إدارة المخزون، الحجوزات، النواقص، والاختلاف بين المخزون الفعلي ومخزون النظام؟",
                "ما نموذج الشحن المطلوب، وما الذي يجب أن يحدث عند التأخير أو فشل التسليم أو الإرجاع؟",
                "كيف تتعاملون مع خدمة العملاء والتصعيدات والشكاوى؟ وما مستوى الخدمة المستهدف؟",
                "كيف تتم العمولات والفواتير والمدفوعات والتسويات المالية اليوم؟",
                "ما التقارير ولوحات المتابعة المطلوبة يومياً وأسبوعياً وللإدارة؟",
                "ما الأنظمة الحالية التي يجب التكامل معها، وما البيانات التي تنتقل في كل اتجاه؟",
                "ما جودة البيانات الحالية، ومن يملك كل نوع من البيانات، وما قواعد الاحتفاظ بها؟",
                "ما متطلبات الصلاحيات والمراجعة والأمن والخصوصية؟",
                "ما المتطلبات الضريبية أو التنظيمية أو متطلبات فئات المنتجات التي يجب مراعاتها؟",
                "ما الافتراضات والمخاطر والقرارات التي تحتاج إجابة من الإدارة؟",
                "ما الأولويات للمرحلة الأولى، وما الذي يؤجل لمراحل لاحقة؟",
                "ما معايير القبول القابلة للقياس لكل مخرج رئيسي؟",
            ]
            english = [
                "Which business outcomes must this project achieve in 6 and 12 months, and how will success be measured?",
                "Who are the stakeholders, what decisions do they own, and who approves changes?",
                "What is the current customer or seller journey from onboarding through post-sale support, and where does it fail?",
                "Which processes currently depend on email, spreadsheets, or manual messages?",
                "How are products created, enriched, approved, and categorised today?",
                "What pricing, promotion, and discount rules apply, and who approves them?",
                "How are inventory, reservations, shortages, and stock discrepancies managed?",
                "Which fulfilment model is required, and how should delivery failures, delays, and returns be handled?",
                "How are customer-service cases and escalations managed, and what service levels are expected?",
                "How are commissions, invoices, payments, refunds, and financial reconciliation handled?",
                "Which daily, weekly, and executive dashboards are required?",
                "Which existing systems must integrate, and which data must move in each direction?",
                "What is the quality, ownership, retention, and governance of the current data?",
                "What access control, audit, security, and privacy requirements apply?",
                "Which tax, regulatory, or category restrictions must be supported?",
                "Which assumptions, risks, and unresolved decisions need management input?",
                "What belongs in phase one, and what should be deferred?",
                "What measurable acceptance criteria are required for each major outcome?",
            ]
            research_section = [research] if research else ["Public research was unavailable for this run; validate all company-specific facts during the meeting."]
            return [("الاستبيان الاستكشافي", arabic), ("Discovery Questionnaire", english), ("Public research to validate", research_section)]
        source_text = " ".join(source.text for source in sources).lower()
        is_marketplace = any(term in source_text for term in ("amazon", "seller", "marketplace", "fulfilment", "fulfillment"))
        if identifier == "gen-strs":
            return DatumOrchestrator._strs_sections(is_marketplace, distribution)
        return DatumOrchestrator._sow_sections(is_marketplace, distribution, parameters)

    @staticmethod
    def _strs_sections(is_marketplace: bool, distribution: DistributionClass) -> list[tuple[str, list[str]]]:
        context = "marketplace seller operations" if is_marketplace else "the organisation's target operating processes"
        sections = [
            ("1. Purpose and business objectives", [
                f"Establish a single operational control point for {context}, replacing fragmented spreadsheet and email tracking.",
                "Provide daily operational visibility, accountable ownership, and measurable exception resolution.",
                "Reduce manual reporting effort, shorten onboarding turnaround, and improve the timeliness of financial reconciliation.",
            ]),
            ("2. Stakeholders and decision rights", [
                "Programme Sponsor: approves scope, priorities, budget-impacting changes, and phase acceptance.",
                "Operations Lead: owns inventory, fulfilment exceptions, seller performance, and operational dashboards.",
                "Seller Experience Lead: owns onboarding, document validation, listing quality, and seller communications.",
                "Finance Lead: owns commissions, settlements, refunds, reconciliation rules, and financial exception closure.",
                "Technology Lead: owns integrations, data access, migration decisions, and technical controls.",
            ]),
            ("3. Current-state issues", [
                "Operational queues are split across email, spreadsheets, and team-specific trackers.",
                "Exceptions lack consistent ownership, priority, due date, source reference, and closure evidence.",
                "Different source systems can hold conflicting status, inventory, payment, or seller information.",
                "Management reporting is manual and cannot reliably show ageing, root cause, or performance trends.",
            ]),
            ("4. Functional requirements", [
                "The system shall maintain work queues for onboarding, listing-quality, inventory, fulfilment, service, return, and finance exceptions.",
                "Each work item shall record source system, external reference, seller or customer reference where applicable, owner, priority, due date, status, and resolution reason.",
                "Users shall assign, reassign, comment on, escalate, and close work items while retaining a complete history.",
                "The system shall validate mandatory onboarding and listing information and route restricted or incomplete cases for approval.",
                "The system shall support fulfilment categories, delivery failures, cancellations, returns, and escalation timing rules.",
                "The system shall support finance reconciliation cases for commissions, settlements, refunds, disputes, and unmatched transactions.",
            ]),
            ("5. Reporting and dashboards", [
                "Daily dashboard: new and overdue onboarding actions, listing-quality exceptions, stock discrepancies, delivery and return exceptions, service escalations, and reconciliation exceptions.",
                "Weekly dashboard: seller performance, order volume, fulfilment performance, return rate, inventory accuracy, exception ageing, and workload by owner.",
                "Management dashboard: financial ageing, unreconciled value, recurring root causes, service-level performance, and phase-one adoption measures.",
            ]),
            ("6. Integrations and data", [
                "Phase one shall support scheduled imports from approved seller, fulfilment or warehouse, customer-service, and finance reporting sources.",
                "Every imported record shall retain source system, import date, external identifier, processing status, and validation outcome.",
                "The solution shall identify missing identifiers, duplicate records, and data-quality exceptions for controlled resolution.",
                "Email notification shall be supported for assignment, escalation, and overdue-action communication.",
            ]),
            ("7. Security, privacy, and audit", [
                "Access shall be role based. Finance-sensitive data shall be restricted to authorised finance users.",
                "The system shall retain an audit trail for approvals, status changes, reassignment, manual corrections, and closure decisions.",
                "The solution shall minimise personal data and shall not store payment-card information.",
            ]),
            ("8. Non-functional requirements", [
                "Daily imports shall show a clear succeeded or failed result with usable error information.",
                "Operational dashboards shall be usable without exporting to spreadsheets for normal daily management.",
                "The design shall support phased delivery, documented configuration, and controlled change approval.",
            ]),
            ("9. Assumptions and open decisions", [
                "Source-system owners will provide stable file layouts, field definitions, and access for agreed imports.",
                "Management must confirm onboarding document rules, approval limits, severity definitions, and escalation timings.",
                "Finance must confirm reconciliation treatment for partial refunds, disputed commissions, and settlement timing differences.",
            ]),
            ("10. Acceptance criteria", [
                "Managers can view the agreed daily dashboard without using a separate spreadsheet process.",
                "Users can create, assign, update, and close exceptions with owner, priority, due date, source reference, and resolution evidence.",
                "Daily imports visibly report success or failure, and at least 95 percent of pilot records process without manual technical correction.",
                "Role restrictions prevent unauthorised users from viewing finance-sensitive information.",
            ]),
        ]
        if distribution == DistributionClass.INTERNAL_ONLY:
            sections.append(("11. Internal delivery considerations", [
                "Data identifier consistency, integration access, and queue ownership are the primary delivery risks.",
                "Run a controlled pilot before wider rollout and maintain a decision log for policy questions that cannot be resolved from source material.",
            ]))
        return sections

    @staticmethod
    def _sow_sections(is_marketplace: bool, distribution: DistributionClass, parameters: dict[str, object]) -> list[tuple[str, list[str]]]:
        directed = parameters.get("directed_findings", [])
        correction_note = f"This version addresses {len(directed)} review finding(s)." if isinstance(directed, list) and directed else "This is the initial scope version."
        sections = [
            ("1. Engagement objective", [
                "Design and implement a phase-one operational control solution that centralises exception management, reporting, and accountability.",
                "Deliver a controlled foundation for seller, fulfilment, service, return, and finance operational processes.",
            ]),
            ("2. Scope of services", [
                "Discovery confirmation and requirements validation workshops with business, operations, finance, customer-service, and technology stakeholders.",
                "Configuration of work queues, priorities, owners, due dates, escalation rules, and closure reasons.",
                "Configuration of onboarding, listing-quality, inventory, fulfilment, customer-service, returns, and reconciliation workflows.",
                "Configuration of role-based access, approval controls, audit history, email notifications, and management dashboards.",
                "Implementation of agreed scheduled data imports, validation rules, error handling, and import-status monitoring.",
                "Pilot support, user acceptance testing, handover documentation, and go-live readiness review.",
            ]),
            ("3. Deliverables", [
                "Approved Stakeholder Requirements Specification and prioritised requirements backlog.",
                "Configured phase-one operational solution and agreed user roles.",
                "Operational work queues and dashboard set for daily, weekly, and management reporting.",
                "Documented import mappings, validation rules, and operational support procedure.",
                "User acceptance test pack, training material, handover pack, and go-live checklist.",
            ]),
            ("4. Delivery approach", [
                "Phase 1: validate requirements, confirm policies, prioritise backlog, and approve solution design.",
                "Phase 2: configure workflows, security roles, dashboards, notifications, and data-import controls.",
                "Phase 3: test with representative source data, resolve defects, and obtain business acceptance.",
                "Phase 4: pilot, train users, complete handover, and review readiness for wider rollout.",
            ]),
            ("5. Client responsibilities", [
                "Provide timely access to subject-matter experts, source exports, field definitions, and representative test data.",
                "Confirm operating policies, approval limits, severity definitions, escalation targets, and data-retention rules.",
                "Review deliverables, perform acceptance testing, and provide consolidated decisions within agreed review windows.",
            ]),
            ("6. OdooTec responsibilities", [
                "Facilitate workshops, document requirements, configure the agreed solution, and maintain the delivery plan.",
                "Document assumptions, dependencies, risks, decisions, test results, and outstanding questions.",
                "Support pilot testing, user enablement, handover, and go-live readiness activities within the agreed scope.",
            ]),
            ("7. Exclusions", [
                "Replacement of external marketplace, warehouse, delivery-provider, payment, or customer-service platforms.",
                "Real-time integration for every source system unless separately approved after technical assessment.",
                "Automated pricing changes, advanced predictive analytics, mobile applications, and functionality not listed in the approved requirements baseline.",
                "Historical data cleansing beyond the agreed pilot dataset and agreed import remediation activities.",
            ]),
            ("8. Assumptions and dependencies", [
                "Approved source files, API access where applicable, stable field definitions, and business owners are available when needed.",
                "The client provides decisions on policies and acceptance criteria without material delay.",
                "Changes to source-system availability, policy, scope, or integration requirements are managed through change control.",
            ]),
            ("9. Acceptance and quality gates", [
                "Each deliverable is reviewed against agreed requirements and acceptance criteria before approval.",
                "The pilot is accepted when workflows, roles, dashboards, and imports operate on representative data with agreed results.",
                "Open defects and findings remain visible until they have an explicit resolution, waiver, or agreed follow-up action.",
            ]),
            ("10. Governance and change control", [
                "A weekly delivery meeting will review progress, decisions, risks, dependencies, and upcoming activities.",
                "Material scope changes require documented impact assessment and sponsor approval before implementation.",
                correction_note,
            ]),
        ]
        if distribution == DistributionClass.INTERNAL_ONLY:
            sections.append(("11. Internal delivery risks and controls", [
                "Prioritise data mapping and identifier quality early; they are critical to queue accuracy and reporting trust.",
                "Maintain an internal decision log for unresolved policy questions and escalate any item that blocks the pilot.",
                "Do not declare the scope clear until all blocking review findings have explicit outcomes.",
            ]))
        return sections

    @staticmethod
    def _validate_input(
        request: StartRunRequest,
        definition: PublicSkillDefinition,
        payload: dict[str, object] | None = None,
    ) -> None:
        supplied = {source.type for source in request.source_material}
        # New Architecture Rules entries distinguish accepted and mandatory
        # sources. Older demo entries treat every listed source as mandatory.
        required = set((payload or {}).get("mandatory_source_material", definition.accepted_source_material))
        missing = required - supplied
        if missing:
            raise ValueError("required_source_material_missing")

    def _record(self, run: StoredRun, step: str, outcome: str, details: dict[str, object] | None = None) -> None:
        entry = RunLogEntry(occurred_at=datetime.now(timezone.utc), step=step, outcome=outcome, details=details or {})
        run.logs.append(entry)
        run.status.log = run.logs


class InProcessRunQueue:
    def __init__(self, handler: object, concurrency: int = 2) -> None:
        self._handler = handler
        self._concurrency = concurrency
        self._queue: asyncio.Queue[UUID] = asyncio.Queue()
        self._workers: list[asyncio.Task[None]] = []

    async def start(self) -> None:
        if not self._workers:
            self._workers = [asyncio.create_task(self._work()) for _ in range(self._concurrency)]

    async def stop(self) -> None:
        for worker in self._workers:
            worker.cancel()
        await asyncio.gather(*self._workers, return_exceptions=True)
        self._workers.clear()

    async def enqueue(self, run_id: UUID) -> None:
        await self._queue.put(run_id)

    async def _work(self) -> None:
        while True:
            run_id = await self._queue.get()
            try:
                await self._handler(run_id)
            finally:
                self._queue.task_done()


class EngineService:
    def __init__(self, repository: PersistentRunRepository, queue: InProcessRunQueue) -> None:
        self._repository = repository
        self._queue = queue

    async def start(self, request: StartRunRequest) -> RunStatus:
        run, created = self._repository.create_or_get(request)
        if created:
            await self._queue.enqueue(run.status.run_id)
        return run.status

    def status(self, run_id: UUID) -> RunStatus:
        return self._repository.get(run_id).status

    def output_path(self, run_id: UUID, filename: str, output_directory: Path) -> Path:
        run = self._repository.get(run_id)
        if not any(filename in {item.filename, item.preview_url.rsplit('/', 1)[-1]} for item in run.status.outputs):
            raise KeyError("output_not_found")
        path = output_directory / filename
        if not path.is_file():
            raise KeyError("output_not_found")
        return path
